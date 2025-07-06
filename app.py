
import streamlit as st
import langchain
from langchain.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.chat_models import ChatOpenAI
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from langchain_groq import ChatGroq
from langchain_mistral import ChatMistral
import os
import os
from dotenv import load_dotenv
load_dotenv() 

os.environ["HF_TOKEN"] =os.getenv("HF_TOKEN")
os.environ["GROQ_API_KEY"] = os.getenv("GROQ_API_KEY")
os.environ["Mistral_API_KEY"] = os.getenv("MISTRAL_API_KEY")

#os.environ["langchain_api_key"] = os.getenv("LANGCHAIN_API_KEY")

from docx import Document
import tempfile

# Step 1: UI - Upload PDF
st.title("📄 PDF Chatbot with RAG")
uploaded_file = st.file_uploader("Upload a PDF", type=["pdf"])

if uploaded_file:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        tmp_file.write(uploaded_file.read())
        pdf_path = tmp_file.name

    # Step 2: Load and Chunk PDF
    loader = PyPDFLoader(pdf_path)
    pages = loader.load_and_split()

    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
    chunks = splitter.split_documents(pages)

    # Step 3: Embed and Create Vector Store (FAISS for local)
    embedding = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    vectorstore = FAISS.from_documents(chunks, embedding)

    # Step 4: Prompt Template
    template = """You are a helpful assistant. Use the context to answer the question.
    Context: {context}
    Question: {question}
    Answer:"""
    prompt = PromptTemplate(input_variables=["context", "question"], template=template)

    # Step 5: QA Chain
    retriever = vectorstore.as_retriever(search_type="mmr", search_kwargs={"k": 5})
    llm = ChatGroq(model="qwen-qwq-32b")
   # llm= ChatMistral(model="mistral-large-v0.1", temperature=0.2, max_tokens=512)
    qa_chain = RetrievalQA.from_chain_type(llm=llm, retriever=retriever, chain_type_kwargs={"prompt": prompt})

    # Step 6: Ask a question
    question = st.text_input("Ask a question about the uploaded PDF:")
    if question:
        with st.spinner("Searching and generating response..."):
            answer = qa_chain.run(question)
        st.success("Answer:")
        st.write(answer)

        # Step 7: Download as DOCX
        doc = Document()
        doc.add_heading("Answer to your question", 0)
        doc.add_paragraph(answer)
        doc_path = os.path.join(tempfile.gettempdir(), "answer.docx")
        doc.save(doc_path)

        with open(doc_path, "rb") as f:
            st.download_button("📥 Download answer as DOCX", f, file_name="response.docx")

